import pikepdf
import tkinter as tk
from tkinter import filedialog, messagebox
import sys
from docx import Document
from lxml import etree
import pdfrw
from pdfrw import PdfReader, PdfWriter, PdfDict, PdfName, PdfString, PdfObject
from datetime import datetime


def extract_data_from_docx(docx_path):
    doc = Document(docx_path)
    xml_content = doc._element.xml
    data = {}

    tree = etree.fromstring(xml_content)
    checkboxes = {}
    foundCheckbox = False
    value = None
    checkboxCounter = 0
    checkboxNamesOrdered = ["Allergien_Nein",
                            "Allergien_Ja",
                            "Unfälle_Nein",
                            "Unfälle_Ja",
                            "Fieber_Nein",
                            "Fieber_Ja",
                            "Beinvenen_Nein",
                            "Beinvenen_Ja",
                            "Frisch_Herzinfarkt_Nein",
                            "Frisch_Herzinfarkt_Ja",
                            "Diabetes_Nein",
                            "Diabetes_Ja",
                            "Bluthochdruck_Nein",
                            "Bluthochdruck_Ja",
                            "Sonst_Krank_Nein",
                            "Sonst_Krank_Ja",
                            "Substanz_Nein",
                            "Substanz_Ja",
                            "Nichtraucher_Nein",
                            "Nichtraucher_Ja",
                            "Raucher_Nein",
                            "Raucher_Ja",
                            "Raucher_Bis20",
                            "Raucher_Über20",
                            "Rauch_Sonst_Nein",
                            "Rauch_Sonst_Ja",
                            "ZNS/PNS_Nein",
                            "ZNS/PNS_Ja",
                            "Anfallsleiden_Nein",
                            "Anfallsleiden_Ja",
                            "Platzangst_Nein",
                            "Platzangst_Ja",
                            "Panikstörung_Nein",
                            "Panikstörung_Ja",
                            "Höhenkrank_Nein",
                            "Höhenkrank_Ja",
                            "Schwindel_Nein",
                            "Schwindel_Ja",
                            "Kopfweh_Nein",
                            "Kopfweh_Ja",
                            "Schlaganfall_Nein",
                            "Schlaganfall_Ja",
                            "Hörgerät_Nein",
                            "Hörgerät_Ja",
                            "Augenkrank_Nein",
                            "Augenkrank_Ja",
                            "Fehlsicht_Nein",
                            "Fehlsicht_Ja",
                            "Kontaktlinsen_Nein",
                            "Kontaktlinsen_Ja",
                            "Brille_Nein",
                            "Brille_Ja",
                            "Netzhautkrank_Nein",
                            "Netzhautkrank_Ja",
                            "Sehnervenkrank_Nein",
                            "Sehnervenkrank_Ja",
                            "Bildfeldausfälle_Nein",
                            "Bildfeldausfälle_Ja",
                            "Husten_Nein",
                            "Husten_Ja",
                            "Husten_Häufig",
                            "Husten_Selten",
                            "Auswurf_Nein",
                            "Auswurf_Ja",
                            "Auswurf_Häufig",
                            "Auswurf_Selten",
                            "Atemnot_Nein",
                            "Atemnot_Ja",
                            "Atemnot_Häufig",
                            "Atemnot_Selten",
                            "Asthma_Nein",
                            "Asthma_Ja",
                            "Lunkeninfarkt_Pneumothorax_Nein",
                            "Lunkeninfarkt_Pneumothorax_Ja",
                            "Schmerzen_Druck_Brust",
                            "Schmerzen_Druck_Rücken",
                            "Schmerzen_Druck_Schultern",
                            "Schmerzen_Druck_Bauch",
                            "Schmerzen_Druck_Nein",
                            "Schmerzen_Druck_Ja",
                            "Herzinfarkt_Nein",
                            "Herzinfarkt_Ja",
                            "Herzrhythmusstörung_Nein",
                            "Herzrhythmusstörung_Ja",
                            "Herzrhythmusstörung_Vorhofflimmern",
                            "Herzrhythmusstörung_Extrasystolen",
                            "Herzrhythmusstörung_Bradykardie",
                            "Herzrhythmusstörung_Tachycardie",
                            "Herzklappenfehler_Nein",
                            "Herzklappenfehler_Ja",
                            "Herzschwäche_Nein",
                            "Herzschwäche_Ja",
                            "Herzinnenhautentzündung_Nein",
                            "Herzinnenhautentzündung_Ja",
                            "Herzschrittmacher_Nein",
                            "Herzschrittmacher_Ja",
                            "Defi_Nein",
                            "Defi_Ja",
                            "Wiederbelebung_Nein",
                            "Wiederbelebung_Ja",
                            "Medikamente_Nein",
                            "Medikamente_Ja",
    ]


    checkboxCounterToDictName = {}
    for i in range(1, len(checkboxNamesOrdered) + 1):
        checkboxCounterToDictName[i] = checkboxNamesOrdered[i - 1]
    yes_counter = 0
    no_counter = 0
    for field in tree.iter():
        # If we find text and a checkbox was just found, map it
        if field is not None and field.text is not None and len(field.text.strip()) > 0 and field.text.strip() != "☐" and field.text.strip() != "☒" and foundCheckbox:
            checkboxCounter += 1
            if checkboxCounter in checkboxCounterToDictName:
                checkboxName = checkboxCounterToDictName[checkboxCounter]
            else:
                checkboxName = str(checkboxCounter)
            checkboxes[checkboxName] = value
            foundCheckbox = False  # Reset for next checkbox

        # Check for a checkbox
        if "checked" in field.tag.lower() and "state" not in field.tag.lower():
            if '1' in field.values():
                value = True
            else:
                value = False
            foundCheckbox = True  # Mark that we found a checkbox
    formerror = False
    truefalseerror = False
    missingInfoError = False
    allergienerror = False
    raucherdatumerror = False
    rauchererror = False
    sonstKrankError = False
    herzInsufError = False
    wrong_fields = list()
    missing_fields = list()
    for checkboxName in checkboxes.keys():
        if checkboxName.endswith("_Ja") and checkboxes[checkboxName[:-3] + "_Nein"] == checkboxes[checkboxName]:
            print(f"Achtung: Die beiden Boxen zum Bereich {checkboxName[:-3]} müssen unterschiedliche Werte haben!")
            wrong_fields.append(checkboxName[:-3])
            formerror = True
            truefalseerror = True
    if checkboxes["Nichtraucher_Ja"] and checkboxes["Raucher_Ja"]:
        print(f"Achtung: Eine Person kann nicht gleichzeitig Raucher und Nichtraucher sein!")
        formerror = True
        rauchererror = True
    if checkboxes["Nichtraucher_Nein"] and checkboxes["Raucher_Nein"]:
        print(f"Achtung: Eine Person muss entweder Raucher oder Nichtraucher sein!")
        formerror = True
        rauchererror = True

    print(checkboxes)
    other_data = extract_table_data(docx_path)

    if ((checkboxes["Nichtraucher_Ja"] and len(other_data["Nichtraucher_Datum"].strip()) != 4)) or ((checkboxes["Raucher_Ja"] and len(other_data["Raucher_Datum"].strip()) != 4)):
        formerror = True
        raucherdatumerror = True

    if checkboxes["Fieber_Ja"] and (other_data["Fieber_Datum"] == "[Empty]" or len(other_data["Fieber_Datum"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Fieberhafter Infekt")

    if checkboxes["Medikamente_Ja"] and (other_data["Medikamente"] == "[Empty]" or len(other_data["Medikamente"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Medikamente Ja aber keine angegeben")

    if checkboxes["Medikamente_Nein"] and other_data["Medikamente"] != "[Empty]" and len(other_data["Medikamente"].strip()) > 0:
        formerror = True
        missingInfoError = True
        missing_fields.append("Medikamente Nein aber welche angegeben")

    if checkboxes["Unfälle_Ja"] and (other_data["Unfälle"] == "[Empty]" or len(other_data["Unfälle"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Operationen/Unfälle")

    if checkboxes["Frisch_Herzinfarkt_Ja"] and (other_data["Frischer_Herzinfarkt_Datum"] == "[Empty]" or len(other_data["Frischer_Herzinfarkt_Datum"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Frischer Herzinfarkt")

    if checkboxes["Diabetes_Ja"] and ((other_data["Diabetis_Diät"] == "[Empty]" or len(other_data["Diabetis_Diät"].strip()) == 0) or (other_data["Diabetis_Medikamente"] == "[Empty]" or len(other_data["Diabetis_Medikamente"].strip()) == 0) or (other_data["Diabetis_Insulin"] == "[Empty]" or len(other_data["Diabetis_Insulin"].strip()) == 0)):
        formerror = True
        missingInfoError = True
        missing_fields.append("Diabetes")

    if checkboxes["Schlaganfall_Ja"] and ((other_data["Schlaganfall_Ausfälle"] == "[Empty]" or len(other_data["Schlaganfall_Ausfälle"].strip()) == 0) or (other_data["Schlaganfall_Datum"] == "[Empty]" or len(other_data["Schlaganfall_Datum"].strip()) == 0)):
        formerror = True
        missingInfoError = True
        missing_fields.append("Schlaganfall")

    if checkboxes["Herzinfarkt_Ja"] and (other_data["Herzinfarkt_Datum"] == "[Empty]" or len(other_data["Herzinfarkt_Datum"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Herzinfarkt (nicht frisch)")

    if checkboxes["Herzklappenfehler_Ja"] and ((other_data["Herzklappenfehler_Grad"] == "[Empty]" or len(other_data["Herzklappenfehler_Grad"].strip()) == 0) or (other_data["Herzklappenfehler_Welche"] == "[Empty]" or len(other_data["Herzklappenfehler_Welche"].strip()) == 0)):
        formerror = True
        missingInfoError = True
        missing_fields.append("Herzklappenfehler")

    if checkboxes["Herzinnenhautentzündung_Ja"] and (other_data["Endokarditis_Datum"] == "[Empty]" or len(other_data["Endokarditis_Datum"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Herzinnenhautentzündung")

    if checkboxes["Herzschrittmacher_Ja"] and (other_data["Herzschrittmacher_Implantatpass"] == "[Empty]" or len(other_data["Herzschrittmacher_Implantatpass"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Herzschrittmacher-Implantation")

    if checkboxes["Defi_Ja"] and (other_data["Defi_Implantatpass"] == "[Empty]" or len(other_data["Defi_Implantatpass"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Defibrillator-Implantation")

    if checkboxes["Wiederbelebung_Ja"] and ((other_data["Herzstillstand_Datum"] == "[Empty]" or len(other_data["Herzstillstand_Datum"].strip()) == 0) or (other_data["Herzstillstand_Situation"] == "[Empty]" or len(other_data["Herzstillstand_Situation"].strip()) == 0)):
        formerror = True
        missingInfoError = True
        missing_fields.append("Wiederbelebung nach Herzsillstand")


    if checkboxes["Sonst_Krank_Ja"] and (other_data["Sonstige_Krankheiten"] == "[Empty]" or len(other_data["Sonstige_Krankheiten"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Sonstige Krankheiten")

    if checkboxes["Herzschwäche_Ja"] and (other_data["Herzinsuffizienz_Beschreibung"] == "[Empty]" or len(other_data["Herzinsuffizienz_Beschreibung"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Herzschwäche (Herzinsuffizienz)")

    if checkboxes["Allergien_Ja"] and (other_data["Allergien"] == "[Empty]" or len(other_data["Allergien"].strip()) == 0):
        formerror = True
        missingInfoError = True
        missing_fields.append("Allergien")

    if formerror:
        errormessage = "Formular falsch ausgefüllt! "
        if truefalseerror:
            errormessage += f"Folgende felder haben entweder keine oder 2 Kreuze: {', '.join(wrong_fields)} "
        if rauchererror:
            errormessage += f"Eine Person ist fälschlicherweise entweder sowohl Raucher als auch Nichtraucher oder keines von beiden."
        if raucherdatumerror:
            errormessage += f"Beim Raucher bzw Nichtraucher feld wird eine 4stellige Jahreszahl erwartet."
        if missingInfoError:
            errormessage += f"Bei den folgenden Feldern wurde ja angekreuzt und dann wurden Informationen vergessen (wenn unbekannt/nicht zutreffend bitte 'unbekannt'/'keine' angeben): {', '.join(missing_fields)}"
        print(errormessage)
        messagebox.showerror("Error", errormessage)
        raise Exception(errormessage)

    other_data["Krankheiten"] = ""
    if checkboxes["Allergien_Ja"]:
        other_data["Krankheiten"] += (f"Allergien: {other_data['Allergien']}, ")
    # TODO check
    if checkboxes["Unfälle_Ja"]:
        other_data["Krankheiten"] += (f"Operationen/Unfälle: {other_data['Unfälle']}, ")
    if checkboxes["Fieber_Ja"]:
        other_data["Krankheiten"] += (f"Fieberhafter Infekt (seit {other_data['Fieber_Datum']}), ")
    if checkboxes["Diabetes_Ja"]:
        other_data["Krankheiten"] += (f"Diabetes (Insulin: {other_data['Diabetis_Insulin']}, Medikamente: {other_data['Diabetis_Medikamente']}, Diät: {other_data['Diabetis_Diät']}), ")
    if checkboxes["Frisch_Herzinfarkt_Ja"]:
        other_data["Krankheiten"] += (f"Frischer Herzinfarkt (Datum: {other_data['Frischer_Herzinfarkt_Datum']}, ")
    if checkboxes["Schlaganfall_Ja"]:
        other_data["Krankheiten"] += (f"Schlaganfall (Datum: {other_data['Schlaganfall_Datum']}, Ausfälle: {other_data['Schlaganfall_Ausfälle']}), ")
    if checkboxes["Beinvenen_Ja"]:
        other_data["Krankheiten"] += (f"Akute Beinvenenthrombose, ")
    if checkboxes["Bluthochdruck_Ja"]:
        other_data["Krankheiten"] += (f"Hypertonie, ")
    if checkboxes["Schmerzen_Druck_Ja"]:
        arten = list()
        if checkboxes['Schmerzen_Druck_Brust']:
            arten.append("Brust")
        if checkboxes['Schmerzen_Druck_Rücken']:
            arten.append("Rücken")
        if checkboxes['Schmerzen_Druck_Schultern']:
            arten.append("Schultern")
        if checkboxes['Schmerzen_Druck_Bauch']:
            arten.append("Bauch")
        if len(arten) > 0:
            other_data["Krankheiten"] += (f"Schmerzen/Druck/Atemnot bei Belastung (Region: {', '.join(arten)}), ")
    if checkboxes["Herzinfarkt_Ja"]:
        other_data["Krankheiten"] += (f"Herzinfarkt (Datum: {other_data['Herzinfarkt_Datum']}), ")
    if checkboxes["Herzrhythmusstörung_Ja"]:
        arten = list()
        if checkboxes['Herzrhythmusstörung_Vorhofflimmern']:
            arten.append("Vorhofflimmern")
        if checkboxes['Herzrhythmusstörung_Extrasystolen']:
            arten.append("Extrasystolen")
        if checkboxes['Herzrhythmusstörung_Bradykardie']:
            arten.append("Bradykardie")
        if checkboxes['Herzrhythmusstörung_Tachycardie']:
            arten.append("Tachycardie")
        if len(arten) > 0:
            other_data["Krankheiten"] += (f"Herzrhythmusstörungen ({', '.join(arten)}), ")
        else:
            other_data["Krankheiten"] += (f"Herzrhythmusstörungen, ")
    if checkboxes["Herzklappenfehler_Ja"]:
        other_data["Krankheiten"] += (f"Herzklappenfehler (Art: {other_data['Herzklappenfehler_Welche']} Grad: {other_data['Herzklappenfehler_Grad']}), ")
    if checkboxes["Herzschwäche_Ja"]:
        other_data["Krankheiten"] += (f"Herzinsuffizienz ({other_data['Herzinsuffizienz_Beschreibung']}), ")
    if checkboxes["Herzinnenhautentzündung_Ja"]:
        other_data["Krankheiten"] += (f"Endokarditis (Datum: {other_data['Endokarditis_Datum']}), ")
    if checkboxes["Herzschrittmacher_Ja"]:
        other_data["Krankheiten"] += (f"Herzschrittmacher-Implantation (Implantatpass: {other_data['Herzschrittmacher_Implantatpass']}), ")
    if checkboxes["Defi_Ja"]:
        other_data["Krankheiten"] += (f"Defibrillator-Implantation (Implantatpass: {other_data['Defi_Implantatpass']}), ")
    if checkboxes["Wiederbelebung_Ja"]:
        other_data["Krankheiten"] += (f"Wiederbelebung nach Herzstillstand (Datum: {other_data['Herzstillstand_Datum']}, Situation: {other_data['Herzstillstand_Situation']}), ")

    if checkboxes["Asthma_Ja"]:
        other_data["Krankheiten"] += (f"Asthma, ")
    if checkboxes["Lunkeninfarkt_Pneumothorax_Ja"]:
        other_data["Krankheiten"] += (f"Zustand nach Lungeninfarkt/Pneumothorax, ")
    if checkboxes["Sonst_Krank_Ja"]:
        other_data["Krankheiten"] += (f"{other_data['Sonstige_Krankheiten']}, ")
    if checkboxes["Anfallsleiden_Ja"]:
        other_data["Krankheiten"] += (f"Anfallsleiden, ")
    if checkboxes["Platzangst_Ja"]:
        other_data["Krankheiten"] += (f"Klaustrophobie, ")
    if checkboxes["Panikstörung_Ja"]:
        other_data["Krankheiten"] += (f"Panikstörung, ")
    if other_data["Krankheiten"].endswith(', '):
        other_data["Krankheiten"] = other_data["Krankheiten"][:-2]
    other_data["Sonstige_Beschwerden"] = ""
    if checkboxes["Anfallsleiden_Ja"]:
        other_data["Sonstige_Beschwerden"] += (f"Anfallsleiden, ")
    if checkboxes["Platzangst_Ja"]:
        other_data["Sonstige_Beschwerden"] += (f"Klaustrophobie, ")
    if checkboxes["Panikstörung_Ja"]:
        other_data["Sonstige_Beschwerden"] += (f"Panikstörung, ")
    if other_data["Sonstige_Beschwerden"].endswith(', '):
        other_data["Sonstige_Beschwerden"] = other_data["Sonstige_Beschwerden"][:-2]
    other_data["Sonstige_Befunde"] = "Keine"
    # print(f"{checkboxes | other_data}")
    return checkboxes | other_data


def extract_table_data(docx_path) -> dict:
    doc = Document(docx_path)
    data_indices_to_field_name = {
        (1, 2, 2): "Allergien",
        (1, 4, 2): "Unfälle",
        (2, 2, 3): "Fieber_Datum",
        (2, 5, 3): "Frischer_Herzinfarkt_Datum",
        (3, 2, 4): "Diabetis_Diät",
        (3, 3, 4): "Diabetis_Medikamente",
        (3, 4, 4): "Diabetis_Insulin",
        (3, 7, 4): "Sonstige_Krankheiten",
        (4, 3, 2): "Nichtraucher_Datum",
        (4, 5, 2): "Raucher_Datum",
        (5, 9, 2): "Schlaganfall_Datum",
        (5, 10, 2): "Schlaganfall_Ausfälle",
        (8, 3, 3): "Herzinfarkt_Datum",
        (8, 7, 3): "Herzklappenfehler_Welche",
        (8, 8, 3): "Herzklappenfehler_Grad",
        (8, 11, 3): "Herzinsuffizienz_Beschreibung",
        (8, 13, 3): "Endokarditis_Datum",
        (8, 15, 3): "Herzschrittmacher_Implantatpass",
        (8, 17, 3): "Defi_Implantatpass",
        (8, 19, 3): "Herzstillstand_Datum",
        (8, 20, 3): "Herzstillstand_Situation",

    }
    medikamente = list()
    collected_data = dict()
    for table_index, table in enumerate(doc.tables):
        # print(f"Table {table_index}:")

        for row_index, row in enumerate(table.rows):
            row_data = []
            for cell_index, cell in enumerate(row.cells):
                cell_text = cell.text.strip() if cell.text.strip() else "[Empty]"
                if table_index == 0:
                    if ":" not in cell_text:
                        continue
                    field_name = cell_text.split(":")[0].strip()
                    if field_name == 'SVNR.':
                        field_value = cell_text.split(":")[1].strip() if cell_text.split(":")[1].strip() else "Feld fehlt"
                        if " " in field_value:
                            collected_data["SVNR_4stellig"] = field_value.split(" ")[0]
                            collected_data["SVNR_datum"] = field_value.split(" ")[1]
                        else:
                            collected_data["SVNR_4stellig"] = field_value[:4]
                            collected_data["SVNR_datum"] = field_value[4:]
                    else:
                        field_value = cell_text.split(":")[1].strip() if cell_text.split(":")[1].strip() else "Feld fehlt"
                        collected_data[field_name] = field_value
                elif (table_index, row_index, cell_index) in data_indices_to_field_name.keys():
                    collected_data[data_indices_to_field_name[(table_index, row_index, cell_index)]] = cell_text.strip()
                row_data.append(cell_text)  # Show "[Empty]" for empty cells
            if table_index == 10 and row_index > 0 and row_data != ['[Empty]', '[Empty]', '[Empty]']:
                medikamente.append(row_data[0] + " (" + row_data[1] + ")")
            # print(f"  Row {row_index}: {row_data}")
            # for cell in range(len(row_data)):
            #     print(f"    Cell {cell}: {row_data[cell]}")

        # print("\n" + "-" * 50 + "\n")  # Separator between tables
    if len(collected_data["Nichtraucher_Datum"]) > 4:
        date_items = collected_data["Nichtraucher_Datum"].split('.')
        if len(date_items[-1]) == 4 and (date_items[-1].startswith("20") or date_items[-1].startswith("19")):
            collected_data["Nichtraucher_Datum"] = date_items[-1]
        elif len(date_items[-1]) == 2 and int(date_items[-1]) < (datetime.now().year % 100):
            collected_data["Nichtraucher_Datum"] = "20" + date_items[-1]
        elif len(date_items[-1]) == 2 and int(date_items[-1]) > (datetime.now().year % 100):
            collected_data["Nichtraucher_Datum"] = "19" + date_items[-1]
    if len(collected_data["Raucher_Datum"]) > 4:
        date_items = collected_data["Raucher_Datum"].split('.')
        if len(date_items[-1]) == 4 and (date_items[-1].startswith("20") or date_items[-1].startswith("19")):
            collected_data["Raucher_Datum"] = date_items[-1]
        elif len(date_items[-1]) == 2 and int(date_items[-1]) < (datetime.now().year % 100):
            collected_data["Raucher_Datum"] = "20" + date_items[-1]
        elif len(date_items[-1]) == 2 and int(date_items[-1]) > (datetime.now().year % 100):
            collected_data["Raucher_Datum"] = "19" + date_items[-1]
    collected_data["Medikamente"] = ", ".join(medikamente)
    # print(collected_data)
    return collected_data

def select_files():
    """Select source and target PDF files"""
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    # For Tests:
    # spirometry_source_pdf = "D:/Rieder Markus_2937211076_LUFU 2_17.04.2025 Anmerkungen.pdf"
    # ergometry_source_pdf = "D:/Rieder Markus_2937211076_ERGO_17.04.2025 ANMERKUNGEN.pdf"
    # word_form = "D:/Einwilligung_Belastungsergometrie_VGÜ_Version2.docx"
    # target_pdf = "D:/uf_hitze_atemschutz_gasrettung_grubenwehr_sauerstoff_2.pdf"
    #
    # # Select source PDF
    print("Bitte Spirometrie Report auswählen!")
    spirometry_source_pdf = filedialog.askopenfilename(
        title="Spirometrie Report auswählen",
        filetypes=[("PDF files", "*.pdf")]
    )

    if not spirometry_source_pdf:
        messagebox.showwarning("Warning", "No file selected!")
        return None, None

    print("Bitte Ergometrie Report auswählen!")
    ergometry_source_pdf = filedialog.askopenfilename(
        title="Ergometrie Report auswählen",
        filetypes=[("PDF files", "*.pdf")]
    )

    if not ergometry_source_pdf:
        messagebox.showwarning("Warning", "No file selected!")
        return None, None

    print("Bitte Word Formular auswählen!")
    word_form = filedialog.askopenfilename(
        title="Word Formular auswählen",
        filetypes=[("DOC Files", "*.docx")]
    )

    print("Bitte UF Hitze Formular auswählen!")
    # Select target PDF
    target_pdf = filedialog.askopenfilename(
        title="UF_Hitze Formular auswählen",
        filetypes=[("PDF files", "*.pdf")]
    )

    if not target_pdf:
        messagebox.showwarning("Warning", "No target file selected!")
        return None, None

    return spirometry_source_pdf, ergometry_source_pdf, word_form, target_pdf


def extract_ergometry_data(pdf_path):
    """Extract spirometry data from the source PDF"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()

        # Dictionary to store extracted values
        ergonometry_data = {}

        # Extract data using regex patterns
        # Look for the parameter table in the PDF text
        lines = text.split('\n')
        # Find lines that contain the spirometry parameters
        counter = 0
        for line in lines:
            counter += 1
            if counter == 17:
                ergonometry_data["Ergo_Datum_Uhrzeit"] = line.split(" ")[1] + " " + line.split(" ")[2][:-3]
                ergonometry_data["Zielfrequenz"] = line.split(" ")[4]
            elif counter == 18:
                ergonometry_data["Ziellast"] = line.split(" ")[1]
            elif counter == 19:
                ergonometry_data["Maxlast"] = line.split(" ")[2]
                ergonometry_data["Maxlast_Prozent"] = line.split(" ")[-2][1:-3]
            elif counter == 24:
                # print(f"{line.split(' ')[1][-5:-3]} {line.split(' ')[1][-2:]}")
                ergonometry_data["Maxlast_Sekunden"] = str(int(line.split(" ")[1][-5:-3])*60 + int(line.split(" ")[1][-2:]))
            # print(f"{counter}: {line}")

        # print(ergonometry_data)
        return ergonometry_data

    except Exception as e:
        print(f"Error extracting data from source PDF: {str(e)}")
        return {}
def extract_spirometry_data(pdf_path):
    """Extract spirometry data from the source PDF"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()

        # Dictionary to store extracted values
        spirometry_data = {}

        # Define the mapping between source parameter names and target field names
        parameter_mapping = {
            'FVC': {
                'messwert_field': 'FVC_Messwert',
                'sollwert_field': 'FVC_Sollwert',
                'prozent_field': 'FVC_Prozent'
            },
            'FEV1': {
                'messwert_field': 'FEV1_Messwert',
                'sollwert_field': 'FEV1_Sollwert',
                'prozent_field': 'FEV1_Prozent'
            },
            'FEV1/FVC': {
                'messwert_field': 'FEV1FVC_Messwert',
                'sollwert_field': 'FEV1FVC_Sollwert',
                'prozent_field': 'FEV1FVC_Prozent'
            },
            'FEF50': {  # This corresponds to MEF50 in the target form
                'messwert_field': 'MEF50_Messwert',
                'sollwert_field': 'MEF50_Sollwert',
                'prozent_field': 'MEF50_Prozent'
            }
        }

        # Extract data using regex patterns
        # Look for the parameter table in the PDF text
        lines = text.split('\n')
        previousLine = ""
        # Find lines that contain the spirometry parameters
        currentColumn = "None"
        columnCounter = 0
        counterToParam = {1: "FVC", 2: "FEV1", 3: "FEV1/FVC", 4: "PEF", 5: "FEF50", 6: "FEF25-75"}
        counter = 0
        for line in lines:
            counter += 1
            if counter == 11:
                spirometry_data["Spiro_Datum_Uhrzeit"] = line.split(" ")[-2] + " " + line.split(" ")[-1][:-3]
            # print(f"{counter}: {line}")
            if line.endswith("-Sollw. (LLN)"):
                currentColumn = "Sollwert"
                columnCounter = 0
            elif line.endswith("-VorBester Versuch"):
                columnCounter = 0
                currentColumn = "Bestwert"
            elif line.endswith("%Sollw."):
                columnCounter = 0
                currentColumn = "Sollprozent"
            if currentColumn != "None" and columnCounter > 0 and columnCounter <= 5 and columnCounter != 4:
                parameterMap = parameter_mapping[counterToParam[columnCounter]]
                if currentColumn == "Sollwert":
                    if line.startswith("- "):
                        spirometry_data[parameterMap["sollwert_field"]] = line[2:].split(' (')[0]
                    else:
                        spirometry_data[parameterMap["sollwert_field"]] = line.split(' (')[0]
                elif currentColumn == "Bestwert":
                    spirometry_data[parameterMap["messwert_field"]] = line.split('*')[0]
                elif currentColumn == "Sollprozent":
                    spirometry_data[parameterMap["prozent_field"]] = line.split('%')[0]
            columnCounter += 1
            # for param, fields in parameter_mapping.items():
            #     # Pattern to match parameter lines with values
            #     # Example: "FVC (L) 5,48 (4,45) 4,84 88,3%"
            #     pattern = rf'{re.escape(param)}\s+\([^)]+\)\s+([\d,]+)\s+\([^)]+\)\s+([\d,]+)\s+([\d,]+)%'
            #     match = re.search(pattern, line)
            #
            #     if match:
            #         sollwert = match.group(1).replace(',', '.')
            #         messwert = match.group(2).replace(',', '.')
            #         prozent = match.group(3)
            #
            #         spirometry_data[fields['sollwert_field']] = sollwert
            #         spirometry_data[fields['messwert_field']] = messwert
            #         spirometry_data[fields['prozent_field']] = prozent
            #
            #         print(f"Extracted {param}: Sollwert={sollwert}, Messwert={messwert}, %Soll={prozent}%")

        # Special handling for FEV1% (which appears as FEV1/FVC in the table)
        # Look for FEV1%= value in the text
        # fev1_percent_match = re.search(r'FEV1%= ([\d,]+)%', text)
        # if fev1_percent_match:
        #     fev1_percent = fev1_percent_match.group(1).replace(',', '.')
        #     spirometry_data['FEV1FVC_Messwert'] = fev1_percent
        #     print(f"Extracted FEV1%: {fev1_percent}%")

        return spirometry_data

    except Exception as e:
        print(f"Error extracting data from source PDF: {str(e)}")
        return {}


def fill_target_pdf(input_pdf, spirometry_data, ergometry_data, form_data, output_pdf):
    """Fill the target PDF form with extracted spirometry data"""
    parameter_mapping_spiro = {
            'Text_33': 'Spiro_Datum_Uhrzeit',
            'Text_39': 'FVC_Messwert',
            'Text_40': 'FVC_Sollwert',
            'Text_41': 'FVC_Prozent',
            'Text_42': 'FEV1_Messwert',
            'Text_43': 'FEV1_Sollwert',
            'Text_44': 'FEV1_Prozent',
            'Text_45': 'FEV1FVC_Messwert',
            'Text_46': 'FEV1FVC_Sollwert',
            'Text_47': 'FEV1FVC_Prozent',
            'Text_48': 'MEF50_Messwert',
            'Text_49': 'MEF50_Sollwert',
            'Text_50': 'MEF50_Prozent'
        }
    parameter_mapping_ergo = {
        'Text_36': 'Ergo_Datum_Uhrzeit',
        'ergo_76': 'Zielfrequenz',
        'ergo_77': 'Ziellast',
        'Text_59': 'Maxlast',
        'Text_61': 'Maxlast_Prozent',
        'Text_60': 'Maxlast_Sekunden'
    }
    parameter_mapping_form = {
        'Text_01': ['Datum'],
        'extra_2': ['Datum'],
        'Text_11': ['Nachname', 'Vorname'],
        'extra_3': ['Nachname', 'Vorname'],
        'Text_12': ['PLZ'],
        'extra_99': ['PLZ'],
        'Text_13': ['Ort', 'Straße/Hausnummer'],
        'extra_102': ['Ort', 'Straße/Hausnummer'],
        'Text_14': ['SVNR_4stellig'],
        'extra_4': ['SVNR_4stellig'],
        'Text_15': ['SVNR_datum'],
        'extra_5': ['SVNR_datum'],
        'Text_27': ['Größe'],
        'Text_28': ['Gewicht'],
        'Text_23': ['Krankheiten'],
        'Text_24': ['Medikamente'],
        'Text_25': ['Sonstige_Beschwerden'],
        'Text_33_1': ['Sonstige_Befunde'],

    }
    checkbox_to_field = {
        18: "Gravidität_Nein",
        19: "Gravidität_Ja",
        20: "Substanz_Nein",
        21: "Substanz_Ja",
        22: "Arztl_Behand_Nein",
        23: "Arztl_Behand_Ja",
        24: "Raucher_Ja",
        25: "Raucher_Bis20",
        26: "Raucher_Über20",
        27: "Nichtraucher_Ja",
        28: "Medikamente_Nein",
        29: "Medikamente_Ja",
        30: "Kopfweh_Nein",
        31: "Kopfweh_Ja",
        32: "Kopf_oft",
        33: "Husten_Nein",
        34: "Husten_Selten",
        35: "Husten_Häufig",
        36: "Schwindel_Nein",
        37: "Schwindel_Ja",
        38: "Schwindel_oft",
        39: "Auswurf_Nein",
        40: "Auswurf_Selten",
        41: "Auswurf_Häufig",
        42: "Schmerzen_Druck_Nein",
        43: "Schmerzen_Druck_Ja",
        44: "Schmerz_Herz_oft",
        45: "Atemnot_Nein",
        46: "Atemnot_Selten",
        47: "Atemnot_Häufig",
        48: "Panikstörung_Nein",
        49: "Panikstörung_Ja",
        50: "Beklemmung_oft",
        51: "Höhenkrank_Nein",
        52: "Höhenkrank_Ja",
        53: "Haut_Nein",
        54: "Haut_2",
        55: "Schilddrüse_Nein",
        56: "Schilddrüse_2",
        57: "Schilddrüse_3",
        58: "Schilddrüse_4",
        59: "Thorax_Nein",
        60: "Thorax_2",
        61: "Atemform_Nein",
        62: "Atemform_2",
        63: "Klopfschall_Nein",
        64: "Klopfschall_2",
        65: "Klopfschall_3",

        66: "Atemgeräusche_Nein",
        67: "Atemgeräusche_2",
        68: "Atemgeräusche_3",
        69: "Atemgeräusche_4",

        70: "Nebengeräusche_Nein",
        71: "Nebengeräusche_2",
        72: "Nebengeräusche_3",
        73: "Nebengeräusche_4",

        74: "ZNS/PNS_Nein",
        75: "ZNS/PNS_Ja",

        76: "ZNS/PNS_Nein",
        77: "ZNS/PNS_Ja",

        78: "Sehen_1",
        79: "Sehen_2",

        80: "Hörgerät_Nein",
        81: "Hörgerät_Ja",

        82: "Herztöne_Nein",
        83: "Herztöne_2",

        84: "Herzrhythmusstörung_Nein",
        85: "Herzrhythmusstörung_Ja",

        86: "Puls_Ruhe_1",
        87: "Puls_Ruhe_2",

        88: "Puls_Ergo_1",
        89: "Puls_Ergo_2",
    }
    checkbox_counter = 0
    text26Age = -1000
    try:
        with pikepdf.open(input_pdf) as pdf:
            pdf.save(output_pdf)
        template_pdf = PdfReader(output_pdf)
        for page in template_pdf.pages:
            if page.Annots:
                for annotation in page.Annots:
                    field_name = annotation.T and annotation.T[1:-1]  # Extract name
                    if field_name == 'Text_26':
                        text26Age = 0
                    text26Age += 1
                    # print(text26Age)
                    # print(f"{field_name}: {annotation.V} {annotation.AP}")
                    # if field_name and field_name in data:
                    #     print(annotation)
                    #     annotation.V = f"({data[field_name]})"  # Fill text field
                    # update checkbox
                    if field_name and (field_name in ["1", "2", "3", "4", "1_1", "2_2", "01",
                                                      "02"] or "checkbox" in field_name.lower()):
                        checkbox_counter += 1
                        # print(f"Checkbox {checkbox_counter} {field_name}: {annotation.V}")
                        pass
                        # SET A CHECKBOX TO 1
                        if checkbox_counter not in checkbox_to_field:
                            pass
                        elif checkbox_to_field[checkbox_counter] in form_data and form_data[checkbox_to_field[checkbox_counter]]:
                            annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                        elif checkbox_to_field[checkbox_counter] in form_data and not form_data[checkbox_to_field[checkbox_counter]]:
                            annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))
                        elif checkbox_to_field[checkbox_counter] == "Arztl_Behand_Ja":
                            if form_data["Medikamente_Ja"]:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                            else:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))
                        elif checkbox_to_field[checkbox_counter] == "Arztl_Behand_Nein" and True:
                            if form_data["Medikamente_Ja"]:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))
                            else:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                        elif checkbox_to_field[checkbox_counter] == "Sehen_1":
                            if form_data["Augenkrank_Nein"] and form_data["Netzhautkrank_Nein"] and form_data["Sehnervenkrank_Nein"] and \
                                form_data["Bildfeldausfälle_Nein"]:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                            else:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))
                        elif checkbox_to_field[checkbox_counter] == "Sehen_2":
                            if form_data["Augenkrank_Ja"] or \
                                form_data["Netzhautkrank_Ja"] or form_data["Sehnervenkrank_Ja"] or \
                                form_data["Bildfeldausfälle_Ja"]:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                            else:
                                annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                                annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))
                        elif checkbox_to_field[checkbox_counter].endswith("_Nein"):
                            annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('1')))
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('1')))
                        else:
                            annotation.update(pdfrw.PdfDict(V=pdfrw.PdfName('0')))
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('0')))

                    # update free field
                    if field_name in parameter_mapping_spiro:
                        # print(annotation)
                        annotation.V = pdfrw.PdfString(f"({spirometry_data[parameter_mapping_spiro[field_name]]})")
                        annotation.AS = pdfrw.PdfString(f"({spirometry_data[parameter_mapping_spiro[field_name]]})")
                        # annotation.AP = pdfrw.PdfDict()
                    if field_name in parameter_mapping_ergo or f"ergo_{text26Age}" in parameter_mapping_ergo:
                        if f"ergo_{text26Age}" in parameter_mapping_ergo:
                            field_name = f"ergo_{text26Age}"
                        # print(annotation)
                        annotation.V = pdfrw.PdfString(f"({ergometry_data[parameter_mapping_ergo[field_name]]})")
                        annotation.AS = pdfrw.PdfString(f"({ergometry_data[parameter_mapping_ergo[field_name]]})")
                    if field_name in parameter_mapping_form:

                        fieldValue = form_data[parameter_mapping_form[field_name][0]]
                        if len(parameter_mapping_form[field_name]) > 1:
                            for i in range(1, len(parameter_mapping_form[field_name])):
                                fieldValue += f", {form_data[parameter_mapping_form[field_name][i]]}"
                        # print(f"{annotation} V: {annotation.V} AP: {annotation.AP}")
                        annotation.V = pdfrw.PdfString(f"({fieldValue})")
                        annotation.AS = pdfrw.PdfString(f"({fieldValue})")
                    if field_name == "Text_22":
                        if form_data["Raucher_Ja"]:
                            annotation.V = pdfrw.PdfString(f"({form_data['Raucher_Datum']})")
                            annotation.AS = pdfrw.PdfString(f"({form_data['Raucher_Datum']})")
                        elif form_data["Nichtraucher_Ja"]:
                            annotation.V = pdfrw.PdfString(f"({form_data['Nichtraucher_Datum']})")
                            annotation.AS = pdfrw.PdfString(f"({form_data['Nichtraucher_Datum']})")
                    if f"extra_{text26Age}" in parameter_mapping_form:
                        field_name = f"extra_{text26Age}"
                        fieldValue = form_data[parameter_mapping_form[field_name][0]]
                        if len(parameter_mapping_form[field_name]) > 1:
                            for i in range(1, len(parameter_mapping_form[field_name])):
                                fieldValue += f", {form_data[parameter_mapping_form[field_name][i]]}"
                        # print(f"{annotation} V: {annotation.V} AP: {annotation.AP}")
                        annotation.V = pdfrw.PdfString(f"({fieldValue})")
                        annotation.AS = pdfrw.PdfString(f"({fieldValue})")
                        annotation.AP = pdfrw.PdfString(f"({fieldValue})")
                        # annotation.AP = pdfrw.PdfDict()
                    # if text26Age >= 70:
                    #     annotation.V = pdfrw.PdfString(f"(number {text26Age})")

        if template_pdf.Root and template_pdf.Root.AcroForm:
            template_pdf.Root.AcroForm.NeedAppearances = PdfObject('true')
        else:
            print("Warning: PDF does not contain an AcroForm dictionary. Field updates may not be applied or visible.")
        # template_pdf.Root.AcroForm.update(pdfrw.PdfDict(NeedAppearances=pdfrw.PdfObject('true')))

        PdfWriter().write(output_pdf, template_pdf)

        print(f"Successfully created filled PDF: {output_pdf}")
        return True

    except Exception as e:
        print(f"Error filling target PDF: {str(e)}")
        return False

def main():
    """Main function to orchestrate the PDF transfer process"""
    print("PDF Spirometry Data Transfer Tool")
    print("=" * 40)

    # Select files
    spirometry_source_pdf, ergometry_source_pdf, word_form, target_pdf = select_files()

    if not spirometry_source_pdf or not target_pdf or not ergometry_source_pdf:
        print("File selection cancelled or incomplete.")
        return

    print(f"Spirometrie PDF: {spirometry_source_pdf}")
    print(f"Ergometrie PDF: {ergometry_source_pdf}")
    print(f"Word Formular: {word_form}")
    print(f"Target PDF: {target_pdf}")
    print()

    # Extract spirometry data from source PDF
    print("Extracting spirometry data from source PDF...")
    spirometry_data = extract_spirometry_data(spirometry_source_pdf)

    if not spirometry_data:
        print("No spirometry data could be extracted from the source PDF.")
        messagebox.showerror("Error", "No spirometry data found in source PDF!")
        return

    # Extract spirometry data from source PDF
    print("Extracting ergometry data from source PDF...")
    ergometry_data = extract_ergometry_data(ergometry_source_pdf)

    if not ergometry_data:
        print("No ergometry data could be extracted from the source PDF.")
        messagebox.showerror("Error", "No ergometry data found in source PDF!")
        return

    # Extract form data from word document
    print("Extracting form data from word Document...")
    form_data = extract_data_from_docx(word_form)

    if not form_data:
        print("No form data could be extracted from the word Document.")
        messagebox.showerror("Error", "No form data found in word document!")
        return

    # print(f"Extracted {len(spirometry_data)} data points:")
    # for field, value in spirometry_data.items():
    #     print(f"  {field}: {value}")
    # print()

    # Generate output filename
    output_path = target_pdf.replace('.pdf', '_filled.pdf')

    # Fill target PDF
    print("Filling target PDF with extracted data...")
    success = fill_target_pdf(target_pdf, spirometry_data, ergometry_data, form_data, output_path)

    if success:
        messagebox.showinfo("Success", f"PDF successfully filled and saved as:\n{output_path}")
    else:
        messagebox.showerror("Error", "Failed to fill the target PDF. Check console for details.")


if __name__ == "__main__":
    # Check if required libraries are installed
    try:
        import PyPDF2
        import tkinter
    except ImportError as e:
        print(f"Required library not found: {e}")
        print("Please install required libraries:")
        print("pip install PyPDF2")
        sys.exit(1)

    main()